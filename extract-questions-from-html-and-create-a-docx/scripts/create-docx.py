from docx import Document
from docx.shared import Pt, RGBColor

# Open and get the lines of example.txt
file = open('./tmp output files/example.txt')
lines = file.readlines()

# Create the docx file and writes the title into it
document = Document()
title = document.add_heading('Aqui está o docx', level=0)
title.runs[0].font.size = Pt(18)  # Font size
title.runs[0].font.bold = True  # Bold style
title.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Font color

# Read each line of example.txt
for line in lines:
    # '[title]' is a mark to heading
    if '[heading]' in line:
        heading = line.replace('[heading]', '').strip()
        heading = document.add_heading(heading, level=1)  # Add heading
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    # '[text]' is a mark to enunciate and asks
    if '[text]' in line:
        text = line.replace('[text]', '').strip()        
        paragraph = document.add_paragraph(text)  # Add a paragraph
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)

    # '[alternative]' is a mark to alternatives(A, B, C, D or E)
    if '[alternative]' in line:
        alternative = line.replace('[alternative]', '').strip()
        alternative = alternative.replace('.', ')')  # Switch . to ) because I like more of )
        paragraph = document.add_paragraph(alternative)  # Add a paragraph for alternative
        paragraph.style = 'List'  # List style
        
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Answers for questions
    if '[answer]' in line:        
        answer = line.replace('[answer]', '')
        
        paragraph = document.add_paragraph(answer)
        paragraph.style = 'List'

        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Arial' 
            run.font.color.rgb = RGBColor(0, 0, 0)

# Save the docx
document.save('./output files/example.docx')
print('Docx created successfully!')